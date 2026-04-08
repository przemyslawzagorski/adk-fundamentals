"""
Word document generator tool for the Analyst Assistant.
Converts markdown-structured content into a high-quality .docx file using python-docx.

Features:
  - Real Word tables with styled headers and alternating row colors
  - Inline formatting: **bold**, *italic*, `code`
  - Mermaid diagrams rendered via mermaid.ink API → PNG embedded in docx
  - Block-aware markdown parser (handles tables and code fences spanning multiple lines)
"""

import base64
import io
import logging
import re
from datetime import datetime
from pathlib import Path

import requests
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

logger = logging.getLogger(__name__)

OUTPUT_DIR = Path(__file__).parent.parent / "output"
MERMAID_INK_URL = "https://mermaid.ink/img/{}"
MERMAID_TIMEOUT_SEC = 15

# ---------------------------------------------------------------------------
# Brand palette (Comarch blue)
# ---------------------------------------------------------------------------
_BLUE = (0x1F, 0x49, 0x7D)
_BLUE_LIGHT = (0xD9, 0xE2, 0xF3)
_WHITE = (0xFF, 0xFF, 0xFF)


def _rgb(r: int, g: int, b: int) -> RGBColor:
    return RGBColor(r, g, b)


def _hex(r: int, g: int, b: int) -> str:
    return f"{r:02X}{g:02X}{b:02X}"


# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------

def _set_cell_bg(cell, r: int, g: int, b: int) -> None:
    """Set table cell background color via OOXML shading element."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), _hex(r, g, b))
    tcPr.append(shd)


def _add_inline_text(paragraph, text: str) -> None:
    """
    Append text to a paragraph, splitting on **bold**, *italic*, `code` markers.
    Each segment becomes a separate run with appropriate formatting.
    """
    pattern = re.compile(r"\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`")
    last_end = 0
    for m in pattern.finditer(text):
        if m.start() > last_end:
            paragraph.add_run(text[last_end : m.start()])
        if m.group(1):          # **bold**
            run = paragraph.add_run(m.group(1))
            run.bold = True
        elif m.group(2):        # *italic*
            run = paragraph.add_run(m.group(2))
            run.italic = True
        elif m.group(3):        # `code`
            run = paragraph.add_run(m.group(3))
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        last_end = m.end()
    if last_end < len(text):
        paragraph.add_run(text[last_end:])


def _is_table_line(line: str) -> bool:
    """Return True when a line looks like a markdown table row or separator."""
    s = line.rstrip()
    return s.startswith("|") and s.endswith("|") and len(s) >= 3


# ---------------------------------------------------------------------------
# Block renderers
# ---------------------------------------------------------------------------

def _render_table(doc: Document, table_lines: list) -> None:
    """Parse markdown table lines and create a styled Word table."""
    # Drop separator rows like |---|:---:|
    data_lines = [
        ln for ln in table_lines
        if not re.fullmatch(r"[|:\-\s]+", ln.rstrip())
    ]
    if not data_lines:
        return

    rows = []
    for line in data_lines:
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)

    if not rows:
        return

    num_cols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=num_cols)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for ri, row_data in enumerate(rows):
        is_header = ri == 0
        word_row = tbl.rows[ri]
        for ci in range(num_cols):
            cell = word_row.cells[ci]
            p = cell.paragraphs[0]
            cell_text = row_data[ci] if ci < len(row_data) else ""
            _add_inline_text(p, cell_text)
            if is_header:
                _set_cell_bg(cell, *_BLUE)
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = _rgb(*_WHITE)
            elif ri % 2 == 0:
                _set_cell_bg(cell, *_BLUE_LIGHT)

    doc.add_paragraph()   # spacing after table


def _render_mermaid(doc: Document, code: str) -> None:
    """
    Render a Mermaid diagram via mermaid.ink public API and embed the PNG.
    Falls back to a formatted code block if the API is unavailable.
    """
    try:
        encoded = base64.urlsafe_b64encode(code.encode("utf-8")).decode("ascii")
        url = MERMAID_INK_URL.format(encoded)
        logger.info(f"[DocGen] 🎨 Rendering Mermaid diagram ({len(code)} chars)...")
        resp = requests.get(url, timeout=MERMAID_TIMEOUT_SEC)
        resp.raise_for_status()
        doc.add_picture(io.BytesIO(resp.content), width=Inches(5.8))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        logger.info("[DocGen] ✅ Mermaid diagram embedded successfully")
    except Exception as exc:
        logger.warning(f"[DocGen] ⚠️ Mermaid render failed ({exc}) — using code fallback")
        note = doc.add_paragraph("[Diagram — mermaid.ink unavailable. Source below:]")
        note.runs[0].italic = True
        fb = doc.add_paragraph(code)
        if fb.runs:
            fb.runs[0].font.name = "Courier New"
            fb.runs[0].font.size = Pt(8)
        doc.add_paragraph()


def _add_cover_page(doc: Document, document_type: str, title: str) -> None:
    """Add a professional cover page with type label, title, and generation timestamp."""
    doc.add_paragraph()
    doc.add_paragraph()

    p_type = doc.add_paragraph(document_type)
    p_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_type.runs[0]
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = _rgb(*_BLUE)

    p_title = doc.add_paragraph(title)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.runs[0].font.size = Pt(16)

    doc.add_paragraph()

    p_date = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_date.runs[0].font.size = Pt(10)

    doc.add_page_break()


def _render_markdown(doc: Document, content: str) -> None:
    """
    Block-aware markdown-to-Word renderer.

    Handles (in order of precedence):
      ```mermaid ... ```  → PNG diagram via mermaid.ink
      ```...```           → monospace code block
      |table|lines|       → styled Word table
      # / ## / ###        → Word headings
      - / * / •           → bulleted list
      1. 2. 3.            → numbered list
      **bold** *italic* `code`  → inline run formatting
    """
    lines = content.splitlines()
    i = 0

    while i < len(lines):
        stripped = lines[i].rstrip()

        # ── Mermaid code fence ─────────────────────────────────────────────
        if stripped.strip() == "```mermaid":
            mermaid_lines = []
            i += 1
            while i < len(lines) and lines[i].strip() != "```":
                mermaid_lines.append(lines[i])
                i += 1
            i += 1  # consume closing ```
            _render_mermaid(doc, "\n".join(mermaid_lines))
            continue

        # ── Generic code fence ─────────────────────────────────────────────
        if stripped.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1
            if code_lines:
                p = doc.add_paragraph(style="Normal")
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Courier New"
                run.font.size = Pt(9)
            continue

        # ── Table block ────────────────────────────────────────────────────
        if _is_table_line(stripped):
            table_lines = []
            while i < len(lines) and _is_table_line(lines[i].rstrip()):
                table_lines.append(lines[i].rstrip())
                i += 1
            _render_table(doc, table_lines)
            continue

        # ── Headings ───────────────────────────────────────────────────────
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
            i += 1
            continue
        if stripped.startswith("# "):
            p = doc.add_heading(stripped[2:], level=1)
            if p.runs:
                p.runs[0].font.color.rgb = _rgb(*_BLUE)
            i += 1
            continue

        # ── Bullet list ────────────────────────────────────────────────────
        if re.match(r"^[-*•]\s+", stripped):
            text = re.sub(r"^[-*•]\s+", "", stripped)
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_text(p, text)
            i += 1
            continue

        # ── Numbered list ──────────────────────────────────────────────────
        num_m = re.match(r"^\d+\.\s+", stripped)
        if num_m:
            p = doc.add_paragraph(style="List Number")
            _add_inline_text(p, stripped[num_m.end():])
            i += 1
            continue

        # ── Empty line ─────────────────────────────────────────────────────
        if stripped == "":
            doc.add_paragraph()
            i += 1
            continue

        # ── Regular paragraph (with inline formatting) ─────────────────────
        p = doc.add_paragraph(style="Normal")
        _add_inline_text(p, stripped)
        i += 1


def generate_docx(
    document_type: str,
    title: str,
    content: str,
) -> str:
    """
    Generates a Word (.docx) document from structured markdown content.

    Call this tool after the document draft has been approved (critic loop complete).

    Args:
        document_type: One of "LLD", "HLD", or "TEST_CASES".
        title: Human-readable document title (e.g. "Payment Service LLD").
        content: Full document content in markdown format.

    Returns:
        Absolute path to the saved .docx file, or an error message.
    """
    try:
        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_title = re.sub(r"[^\w\-]", "_", title)[:40]
        filename = f"{document_type}_{safe_title}_{timestamp}.docx"
        output_path = OUTPUT_DIR / filename

        doc = Document()

        # Margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.2)
            section.right_margin = Inches(1.2)

        _add_cover_page(doc, document_type, title)
        _render_markdown(doc, content)

        doc.save(output_path)
        logger.info(f"Document saved: {output_path}")
        return f"✅ Document saved: {output_path}"

    except Exception as e:
        logger.error(f"Error generating document: {e}")
        return f"❌ Error generating document: {e}"

